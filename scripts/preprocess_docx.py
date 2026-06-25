#!/usr/bin/env python3
"""
DOCX Preprocessing Pipeline for SE Requirements Decomposition.

Converts Word documents (PRDs, customer specs, industry standards, meeting notes)
into structured Markdown — preserving headings, paragraphs, tables, and embedded images.

Usage:
    python preprocess_docx.py --input PRD.docx --output ./preprocessed/
    python preprocess_docx.py --input ./docs/inputs/ --output ./preprocessed/ --batch
"""

import argparse
import json
import logging
import os
import sys
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

logger = logging.getLogger("docx_preprocess")
logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class SectionInfo:
    """Metadata for one document section (heading level + content)."""
    level: int  # 0 = document title, 1 = H1, 2 = H2, 3 = H3
    heading: str
    paragraph_count: int = 0
    table_count: int = 0
    image_count: int = 0


@dataclass
class DocxManifest:
    """Output manifest for one Word document."""
    source: str
    total_paragraphs: int
    total_tables: int
    total_images: int
    sections: list[SectionInfo] = field(default_factory=list)
    issues: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Core extraction
# ---------------------------------------------------------------------------

def _load_docx(filepath: str):
    """Load a .docx file using python-docx. Returns (doc, issues)."""
    try:
        from docx import Document
        doc = Document(filepath)
        return doc, []
    except ImportError:
        return None, ["python-docx not installed. Run: pip install python-docx"]
    except Exception as exc:
        return None, [f"Failed to open {filepath}: {exc}"]


def _is_heading(paragraph) -> bool:
    """Check if a paragraph is a heading (built-in style or outline level)."""
    style_name = (paragraph.style.name if paragraph.style else "").lower()
    if any(kw in style_name for kw in ("heading", "heading 1", "heading 2",
                                         "heading 3", "title", "subtitle")):
        return True
    # Check outline level (0-3 = heading in most templates)
    pPr = paragraph._element.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
    if pPr is not None:
        outline = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl")
        if outline is not None:
            level = int(outline.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "9"))
            if level <= 3:
                return True
    return False


def _get_heading_level(paragraph) -> int:
    """Extract numeric heading level (0=title, 1=H1, 2=H2, 3=H3+)."""
    style_name = (paragraph.style.name if paragraph.style else "").lower()
    if "title" in style_name:
        return 0
    for level in (1, 2, 3):
        if f"heading {level}" in style_name:
            return level
    # Check outline level
    pPr = paragraph._element.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
    if pPr is not None:
        outline = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl")
        if outline is not None:
            ol = int(outline.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "9"))
            return ol + 1  # outline level 0 → heading 1
    return 1  # default


def _extract_table_as_markdown(table) -> str:
    """Convert a python-docx Table to GitHub-flavored Markdown."""
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            text = cell.text.replace("\n", " ").replace("|", "\\|").strip()
            cells.append(text)
        rows.append(cells)

    if not rows:
        return ""

    max_cols = max(len(r) for r in rows)
    # Pad short rows
    for r in rows:
        while len(r) < max_cols:
            r.append("")

    # Trim trailing empty columns
    while max_cols > 1 and all(r[-1] == "" for r in rows):
        for r in rows:
            r.pop()
        max_cols -= 1

    lines = [
        "| " + " | ".join(rows[0]) + " |",
        "| " + " | ".join(["---"] * max_cols) + " |",
    ]
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines) + "\n"


def _extract_images(doc, output_dir: str) -> int:
    """Extract embedded images from the docx (ZIP of XML + media)."""
    import zipfile

    count = 0
    images_dir = os.path.join(output_dir, "images")
    os.makedirs(images_dir, exist_ok=True)

    # DOCX is a ZIP — extract image parts directly
    docx_path = None
    if hasattr(doc, 'part') and hasattr(doc.part, 'package'):
        # Try to find the source path
        pass

    # Fallback: user provides the filepath externally
    # We handle this during process_docx() by re-opening as ZIP
    return count


def extract_images_from_docx(filepath: str, output_dir: str) -> int:
    """Extract embedded images from .docx by reading the ZIP directly."""
    import zipfile
    import re
    import io

    images_dir = os.path.join(output_dir, "images")
    os.makedirs(images_dir, exist_ok=True)
    count = 0
    image_exts = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".tif", ".emf", ".wmf", ".svg"}

    try:
        with zipfile.ZipFile(filepath, "r") as z:
            for name in z.namelist():
                ext = Path(name).suffix.lower()
                if ext in image_exts and "media" in name.lower():
                    data = z.read(name)
                    if len(data) < 100:
                        continue  # skip tiny (likely decoration)
                    out_name = Path(name).name
                    out_path = os.path.join(images_dir, out_name)
                    with open(out_path, "wb") as f:
                        f.write(data)
                    logger.info(f"  Image: {out_name} ({len(data)} bytes)")
                    count += 1
    except Exception as exc:
        logger.warning(f"Image extraction via ZIP failed: {exc}")

    if count > 0:
        logger.info(f"Extracted {count} image(s) → {images_dir}")
    return count


# ---------------------------------------------------------------------------
# Quality verification
# ---------------------------------------------------------------------------

def verify_docx_quality(manifest: DocxManifest, text_path: str) -> dict[str, Any]:
    """Assess DOCX extraction quality."""
    checks: list[dict[str, Any]] = []
    score = 100.0

    # Check 1: Paragraph density
    if manifest.total_paragraphs == 0 and manifest.total_tables == 0:
        score -= 50
        checks.append({"dimension": "content", "status": "warn", "detail": "No paragraphs or tables extracted — document may be empty or unreadable"})
    elif manifest.total_paragraphs < 5:
        score -= 15
        checks.append({"dimension": "content", "status": "info", "detail": f"Only {manifest.total_paragraphs} paragraphs — short document or extraction issue"})

    # Check 2: Section hierarchy
    if manifest.sections:
        levels = set(s.level for s in manifest.sections)
        if 1 not in levels:
            checks.append({"dimension": "structure", "status": "info", "detail": "No Heading 1 found — document may lack top-level structure"})
    elif manifest.total_paragraphs > 10:
        checks.append({"dimension": "structure", "status": "warn", "detail": "No headings detected but has paragraphs — possible flat structure or heading detection failure"})

    # Check 3: Table coverage
    if manifest.total_tables > 0:
        empty_tables = sum(1 for s in manifest.sections if s.table_count == 0)
        checks.append({"dimension": "tables", "status": "info", "detail": f"{manifest.total_tables} table(s) across {len(manifest.sections)} sections"})

    # Check 4: Image extraction
    if manifest.total_images > 0:
        checks.append({"dimension": "images", "status": "info", "detail": f"{manifest.total_images} image(s) extracted — verify they are requirements-relevant"})

    # Check 5: Issues
    if manifest.issues:
        score -= len(manifest.issues) * 15
        checks.append({"dimension": "issues", "status": "warn", "detail": f"{len(manifest.issues)} issue(s): {'; '.join(manifest.issues[:3])}"})

    if any(c["status"] == "warn" for c in checks):
        score = min(score, 85)

    return {"overall_score": round(max(score, 0), 1), "checks": checks}


def print_docx_quality(report: dict[str, Any]) -> None:
    icon = "[PASS]" if report["overall_score"] >= 90 else ("[WARN]" if report["overall_score"] >= 70 else "[FAIL]")
    print(f"""
======================================================================
  DOCX Extraction Quality Report
======================================================================
  Overall:    {icon} {report['overall_score']:.0f}/100
----------------------------------------------------------------------""")
    for c in report["checks"]:
        ic = {"pass": "[OK]  ", "warn": "[WARN]", "info": "[INFO]"}.get(c["status"], "      ")
        print(f"  {ic} {c['detail'][:67]}")
    print("======================================================================\n")


# ---------------------------------------------------------------------------
# Main pipeline
# ---------------------------------------------------------------------------

def process_docx(
    filepath: str,
    output_dir: str,
    *,
    extract_images: bool = True,
    verify: bool = False,
) -> DocxManifest:
    """Run full preprocessing on a single .docx file."""
    os.makedirs(output_dir, exist_ok=True)

    # Handle legacy .doc (skip — user should convert first)
    ext = Path(filepath).suffix.lower()
    if ext == ".doc":
        manifest = DocxManifest(
            source=os.path.basename(filepath),
            total_paragraphs=0,
            total_tables=0,
            total_images=0,
            issues=[".doc format not supported — convert to .docx first using LibreOffice: "
                     "soffice --headless --convert-to docx file.doc"],
        )
        _write_manifest(manifest, output_dir)
        return manifest

    doc, issues = _load_docx(filepath)
    if doc is None:
        manifest = DocxManifest(
            source=os.path.basename(filepath),
            total_paragraphs=0,
            total_tables=0,
            total_images=0,
            issues=issues,
        )
        _write_manifest(manifest, output_dir)
        return manifest

    # --- Step 1: Walk document body ---
    md_parts: list[str] = []
    sections: list[SectionInfo] = []
    current_section: SectionInfo | None = None
    para_count = 0
    table_count = 0
    table_idx = 0

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Find the matching paragraph object
            para = _find_paragraph(doc, element)
            if para is None:
                continue

            text = para.text.strip()
            if not text:
                continue

            if _is_heading(para):
                level = _get_heading_level(para)
                prefix = "#" * min(level + 1, 4)  # title → # , H1 → ## , H2 → ###
                md_parts.append(f"\n{prefix} {text}\n")
                current_section = SectionInfo(
                    level=level, heading=text,
                )
                sections.append(current_section)
            else:
                md_parts.append(f"{text}\n")
                para_count += 1
                if current_section:
                    current_section.paragraph_count += 1

        elif tag == "tbl":
            table = _find_table(doc, element)
            if table is not None:
                table_idx += 1
                md = _extract_table_as_markdown(table)
                md_parts.append(f"\n{md}\n")
                table_count += 1
                if current_section:
                    current_section.table_count += 1

    # --- Step 2: Write Markdown ---
    full_md = "\n".join(md_parts)
    md_path = os.path.join(output_dir, "full-text.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(f"<!-- Source: {os.path.basename(filepath)} -->\n\n")
        f.write(full_md)
    logger.info(f"Full text → {md_path} ({para_count} paragraphs, {table_count} tables)")

    # --- Step 3: Extract images ---
    image_count = 0
    if extract_images:
        image_count = extract_images_from_docx(filepath, output_dir)

    # --- Step 4: Build manifest ---
    manifest = DocxManifest(
        source=os.path.basename(filepath),
        total_paragraphs=para_count,
        total_tables=table_count,
        total_images=image_count,
        sections=sections,
        issues=issues,
    )
    _write_manifest(manifest, output_dir)
    _print_summary(manifest)

    if verify:
        text_path = os.path.join(output_dir, "full-text.md")
        report = verify_docx_quality(manifest, text_path)
        report_path = os.path.join(output_dir, "quality-report.json")
        with open(report_path, "w", encoding="utf-8") as f:
            json.dump(report, f, indent=2, ensure_ascii=False)
        print_docx_quality(report)
        logger.info(f"Quality report -> {report_path}")

    return manifest


def _find_paragraph(doc, element):
    """Find the python-docx Paragraph object matching an lxml element."""
    for para in doc.paragraphs:
        if para._element is element:
            return para
    return None


def _find_table(doc, element):
    """Find the python-docx Table object matching an lxml element."""
    for table in doc.tables:
        if table._element is element:
            return table
    return None


def _write_manifest(manifest: DocxManifest, output_dir: str) -> str:
    """Write manifest.json."""
    path = os.path.join(output_dir, "manifest.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(asdict(manifest), f, indent=2, ensure_ascii=False)
    logger.info(f"Manifest → {path}")
    return path


def _print_summary(m: DocxManifest) -> None:
    print(f"""
═══════════════════════════════════════════
  DOCX Preprocessing Complete
═══════════════════════════════════════════
  Source:          {m.source}
  Paragraphs:      {m.total_paragraphs}
  Tables:          {m.total_tables}
  Images:          {m.total_images}
  Sections:        {len(m.sections)}
  Issues:          {len(m.issues)}
═══════════════════════════════════════════
""")


def process_batch(input_dir: str, output_base: str) -> list[DocxManifest]:
    """Process all .docx files in a directory."""
    results: list[DocxManifest] = []
    files = list(Path(input_dir).glob("*.docx"))
    logger.info(f"Batch processing {len(files)} DOCX file(s) in {input_dir}")

    for f in files:
        output_dir = os.path.join(output_base, f.stem)
        logger.info(f"--- Processing: {f.name} ---")
        manifest = process_docx(str(f), output_dir)
        results.append(manifest)

    return results


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="DOCX Preprocessing for SE Requirements Decomposition",
    )
    parser.add_argument(
        "--input", "-i", required=True,
        help="Path to a .docx file or directory (with --batch)",
    )
    parser.add_argument(
        "--output", "-o", required=True,
        help="Output directory for preprocessed artifacts",
    )
    parser.add_argument(
        "--batch", action="store_true",
        help="Process all .docx files in the input directory",
    )
    parser.add_argument(
        "--no-images", action="store_true",
        help="Skip embedded image extraction",
    )
    parser.add_argument(
        "--verify", action="store_true",
        help="Run quality assessment on extraction results",
    )
    parser.add_argument(
        "--verbose", "-v", action="store_true",
        help="Verbose logging",
    )

    args = parser.parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    if args.batch:
        if not os.path.isdir(args.input):
            print(f"ERROR: --batch requires a directory, got: {args.input}", file=sys.stderr)
            sys.exit(1)
        process_batch(args.input, args.output)
    else:
        if not os.path.isfile(args.input):
            print(f"ERROR: File not found: {args.input}", file=sys.stderr)
            sys.exit(1)
        process_docx(args.input, args.output, extract_images=not args.no_images, verify=args.verify)


if __name__ == "__main__":
    main()
